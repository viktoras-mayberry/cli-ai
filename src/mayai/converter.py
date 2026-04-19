"""File format conversion registry."""

from __future__ import annotations

import csv
import json
import os
from pathlib import Path

from .display import print_info, print_error, print_success, print_warning


def _confirm(prompt: str = "Proceed? [y/N] ") -> bool:
    try:
        return input(prompt).strip().lower() in ("y", "yes")
    except (EOFError, KeyboardInterrupt):
        return False


def _human_size(size_bytes: int) -> str:
    for unit in ("B", "KB", "MB", "GB"):
        if abs(size_bytes) < 1024:
            return f"{size_bytes:.1f} {unit}" if unit != "B" else f"{size_bytes} {unit}"
        size_bytes /= 1024  # type: ignore[assignment]
    return f"{size_bytes:.1f} TB"


# ------------------------------------------------------------------ #
# Individual converters                                                #
# ------------------------------------------------------------------ #

def _csv_to_xlsx(src: Path, dst: Path) -> None:
    try:
        from openpyxl import Workbook
    except ImportError:
        raise RuntimeError("CSV-to-Excel requires openpyxl. Install: pip install openpyxl")

    wb = Workbook()
    ws = wb.active
    with open(src, newline="", encoding="utf-8", errors="replace") as f:
        for row in csv.reader(f):
            ws.append(row)
    wb.save(str(dst))


def _xlsx_to_csv(src: Path, dst: Path) -> None:
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise RuntimeError("Excel-to-CSV requires openpyxl. Install: pip install openpyxl")

    wb = load_workbook(str(src), read_only=True, data_only=True)
    ws = wb.active
    with open(dst, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        for row in ws.iter_rows(values_only=True):
            writer.writerow(row)
    wb.close()


def _json_to_csv(src: Path, dst: Path) -> None:
    data = json.loads(src.read_text(encoding="utf-8"))
    if isinstance(data, list) and data and isinstance(data[0], dict):
        keys = list(data[0].keys())
        with open(dst, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=keys)
            writer.writeheader()
            writer.writerows(data)
    else:
        raise RuntimeError("JSON-to-CSV requires a JSON array of objects.")


def _csv_to_json(src: Path, dst: Path) -> None:
    with open(src, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.DictReader(f)
        data = list(reader)
    dst.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def _image_convert(src: Path, dst: Path) -> None:
    try:
        from PIL import Image
    except ImportError:
        raise RuntimeError("Image conversion requires Pillow. Install: pip install Pillow")

    img = Image.open(str(src))
    if img.mode in ("RGBA", "P") and dst.suffix.lower() in (".jpg", ".jpeg"):
        img = img.convert("RGB")
    img.save(str(dst))


def _txt_to_docx(src: Path, dst: Path) -> None:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("Text-to-Word requires python-docx. Install: pip install python-docx")

    doc = Document()
    text = src.read_text(encoding="utf-8", errors="replace")
    for para in text.split("\n"):
        doc.add_paragraph(para)
    doc.save(str(dst))


def _docx_to_txt(src: Path, dst: Path) -> None:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("Word-to-text requires python-docx. Install: pip install python-docx")

    doc = Document(str(src))
    text = "\n".join(p.text for p in doc.paragraphs)
    dst.write_text(text, encoding="utf-8")


# ------------------------------------------------------------------ #
# Registry                                                             #
# ------------------------------------------------------------------ #

_CONVERTERS: dict[tuple[str, str], callable] = {
    ("csv", "xlsx"): _csv_to_xlsx,
    ("xlsx", "csv"): _xlsx_to_csv,
    ("xls", "csv"): _xlsx_to_csv,
    ("json", "csv"): _json_to_csv,
    ("csv", "json"): _csv_to_json,
    ("txt", "docx"): _txt_to_docx,
    ("md", "docx"): _txt_to_docx,
    ("docx", "txt"): _docx_to_txt,
    # Image conversions
    ("png", "jpg"): _image_convert,
    ("png", "jpeg"): _image_convert,
    ("jpg", "png"): _image_convert,
    ("jpeg", "png"): _image_convert,
    ("png", "webp"): _image_convert,
    ("webp", "png"): _image_convert,
    ("jpg", "webp"): _image_convert,
    ("jpeg", "webp"): _image_convert,
    ("webp", "jpg"): _image_convert,
    ("webp", "jpeg"): _image_convert,
    ("bmp", "png"): _image_convert,
    ("bmp", "jpg"): _image_convert,
}


def get_supported_conversions() -> list[str]:
    """Return human-readable list of supported conversions."""
    return [f"{src} -> {dst}" for src, dst in sorted(_CONVERTERS.keys())]


def convert_file(
    source: str,
    target_format: str,
    output_path: str | None = None,
    *,
    auto_confirm: bool = False,
) -> str | None:
    """Convert a file to another format. Returns output path or None.

    If output_path is not specified, generates one by changing the extension.
    """
    src = Path(source)
    if not src.exists():
        print_error(f"File not found: {source}")
        return None

    src_ext = src.suffix.lower().lstrip(".")
    dst_ext = target_format.lower().lstrip(".")

    converter = _CONVERTERS.get((src_ext, dst_ext))
    if converter is None:
        print_error(
            f"Conversion from .{src_ext} to .{dst_ext} is not supported.\n"
            f"Supported: {', '.join(get_supported_conversions())}"
        )
        return None

    if output_path:
        dst = Path(output_path)
    else:
        dst = src.with_suffix(f".{dst_ext}")

    src_size = _human_size(src.stat().st_size)
    print_info(f"Convert: {src.name} ({src_size}) -> {dst.name}")

    if not auto_confirm and not _confirm():
        print_info("Cancelled.")
        return None

    try:
        converter(src, dst)
        dst_size = _human_size(dst.stat().st_size)
        print_success(f"Converted: {dst} ({dst_size})")
        return str(dst)
    except RuntimeError as exc:
        print_error(str(exc))
        return None
    except Exception as exc:
        print_error(f"Conversion failed: {exc}")
        return None
